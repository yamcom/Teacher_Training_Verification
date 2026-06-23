# services/report_service.py
import os
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from models import Teacher, TrainingRecord
from services.tag_service import TagService

class ReportService:
    """研習報表與統計分析服務，負責計算標籤通過率與生成正式分析報告"""

    @classmethod
    def calculate_tag_statistics(cls) -> pd.DataFrame:
        """
        3.6 統計分析功能：依標籤統計全校通過人數、未通過人數與通過率。
        通過率計算公式：通過人數 ÷ 全校教師人數 × 100%
        """
        # 取得全校總教師人數
        total_teachers = Teacher.query.count()
        if total_teachers == 0:
            return pd.DataFrame(columns=['研習類別', '通過人數', '未通過人數', '通過率'])

        # 初始化統計字典
        stats_dict = {tag: {'passed': 0, 'failed': 0} for tag in TagService.PRESET_TAGS}

        # 撈取所有教師的研習紀錄進行分析
        teachers = Teacher.query.all()
        for teacher in teachers:
            # 收集該教師所有「通過(時數>0)」的標籤
            teacher_passed_tags = set()
            passed_records = [r for r in teacher.records if r.hours > 0]
            
            for r in passed_records:
                tags = TagService.parse_training_name(r.training.title)
                for tag in tags:
                    teacher_passed_tags.add(tag)
            
            # 根據該教師的通過狀態，更新全校各標籤的數據
            for tag in TagService.PRESET_TAGS:
                if tag in teacher_passed_tags:
                    stats_dict[tag]['passed'] += 1
                else:
                    stats_dict[tag]['failed'] += 1

        # 整理為 DataFrame 格式
        stats_data = []
        for tag in TagService.PRESET_TAGS:
            passed_count = stats_dict[tag]['passed']
            failed_count = stats_dict[tag]['failed']
            # 計算通過率
            passing_rate = (passed_count / total_teachers) * 100
            
            stats_data.append({
                '研習類別': tag,
                '通過人數': passed_count,
                '未通過人數': failed_count,
                '通過率': f"{passing_rate:.1f}%"
            })

        return pd.DataFrame(stats_data)

    @classmethod
    def export_tag_analysis_report(cls, output_dir: str) -> str:
        """
        第八章 報表輸出：匯出「標籤統計分析表」之正式 Word 報告 (.docx)
        包含精美標題、統計概述與美化過的分析表格。
        """
        file_path = os.path.join(output_dir, '數位學習研習標籤統計分析報告.docx')
        
        doc = Document()
        
        # 1. 設置報告大標題
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run('教師數位學習增能研習指標  統計分析報告')
        title_run.bold = True
        title_run.font.size = Pt(20)
        title_run.font.color.rgb = RGBColor(31, 78, 121) # 深藍色行政風格
        
        # 2. 引言與基本資訊
        meta_p = doc.add_paragraph()
        meta_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        meta_p.add_run('報告產出日期：115年06月23日\n產出單位：教務處處務組')
        
        doc.add_heading('一、 統計目的與概述', level=1)
        desc_p = doc.add_paragraph(
            '本報告依據「推動中小學數位學習精進方案」實施計畫辦理，針對全校教師進行 A1（數位學習工作坊一）、'
            'A2（數位學習工作坊二）及 B 系列（數增能進階課程）之修課紀錄進行大數據智慧檢核。'
            '旨在掌握全校教師研習達成率，作為後續課表調配與行政送件之依據。'
        )
        
        # 3. 獲取動態計算數據
        df_stats = cls.calculate_tag_statistics()
        total_teachers = Teacher.query.count()
        
        summary_p = doc.add_paragraph()
        summary_p.add_run(f'目前全校編制教師總數為：{total_teachers} 人。各項指標分析摘要如下：').bold = True

        # 4. 建立美化版統計表格
        # 欄位：研習類別, 通過人數, 未通過人數, 通過率 (共 4 欄)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # 設定表頭文字與底色
        hdr_cells = table.rows[0].cells
        headers = ['研習類別 (指標)', '已通過人數', '未通過人數', '全校通過率']
        for i, header_text in enumerate(headers):
            hdr_cells[i].text = header_text
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            cls._set_cell_background(hdr_cells[i], "1F4E79") # 表頭深藍底色
            hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255) # 白色字
            hdr_cells[i].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 填入大數據統計列
        for _, row in df_stats.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['研習類別'])
            row_cells[1].text = str(row['通過人數'])
            row_cells[2].text = str(row['未通過人數'])
            row_cells[3].text = str(row['通過率'])
            
            # 文字置中對齊
            for cell in row_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 5. 結論建議
        doc.add_heading('\n二、 後續策進建議', level=1)
        suggest_p = doc.add_paragraph()
        suggest_p.add_run('1. 針對通過率未達 80% 之指標（如 B5-1、B5-2 生成式 AI 應用），教務處將於下學期校務會議期間加開校內集中研習場次。\n')
        suggest_p.add_run('2. 系統已自動篩選出「未通過教師名冊」，後續將由資訊組個別發送 Email 或 LINE 提醒通知補修相關時數。')

        doc.save(file_path)
        return file_path

    @staticmethod
    def _set_cell_background(cell, fill_hex):
        """私有輔助方法：透過 XML 設定 Word 表格單格底色"""
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_hex)
        tc_pr.append(shd)