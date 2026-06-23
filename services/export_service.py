# services/export_service.py
import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from models import Teacher, TrainingRecord
from services.statistics_service import StatisticsService

class ExportService:
    """報表輸出服務，負責產出 Excel、CSV 與正式 Word 公文報表"""

    @classmethod
    def export_to_double_column_csv(cls, file_path: str) -> str:
        """
        將全校研習統計表轉換為符合「附件2樣板」的雙欄 CSV 格式
        (左側一欄、右側一欄，每欄包含：職稱, 姓名, 通過研習)
        """
        # 1. 取得標準單欄的統計結果 (序號, 職位, 姓名, 通過研習)
        df_single = StatisticsService.generate_all_school_report()
        
        # 2. 將資料依原始順序平分成左右兩半
        total_rows = len(df_single)
        mid_point = (total_rows + 1) // 2  # 無條件進位平分
        
        left_half = df_single.iloc[:mid_point].reset_index(drop=True)
        right_half = df_single.iloc[mid_point:].reset_index(drop=True)
        
        # 3. 建立雙欄合併後的 DataFrame
        double_column_data = []
        for i in range(mid_point):
            row_data = {}
            
            # 左半邊資料
            row_data['職稱_左'] = left_half.loc[i, '職位'] if i < len(left_half) else ''
            row_data['姓名_左'] = left_half.loc[i, '姓名'] if i < len(left_half) else ''
            row_data['通過研習_左'] = left_half.loc[i, '通過研習'] if i < len(left_half) else ''
            
            # 右半邊資料
            row_data['職稱_右'] = right_half.loc[i, '職位'] if i < len(right_half) else ''
            row_data['姓名_右'] = right_half.loc[i, '姓名'] if i < len(right_half) else ''
            row_data['通過研習_右'] = right_half.loc[i, '通過研習'] if i < len(right_half) else ''
            
            double_column_data.append(row_data)
            
        df_double = pd.DataFrame(double_column_data)
        
        # 4. 設定符合主管機關審查用的欄位標題
        df_double.columns = ['職稱', '姓名', '通過研習', '職稱', '姓名', '通過研習']
        
        # 5. 轉換為帶有 UTF-8 BOM 的 CSV 字串 (防止 Excel 開啟亂碼)
        return df_double.to_csv(index=False, encoding='utf-8-sig')

    @classmethod
    def export_to_excel(cls, output_dir: str) -> str:
        """
        將全校研習統計表與詳細明細清單分開輸出至同一個 Excel 活頁簿 (.xlsx) 中
        """
        file_path = os.path.join(output_dir, '全校教師研習統計總表.xlsx')
        
        # 取得統計總表
        df_summary = StatisticsService.generate_all_school_report()
        
        # 取得所有通過的研習詳細明細清單
        teachers = Teacher.query.all()
        detail_data = []
        for teacher in teachers:
            for record in teacher.records:
                if record.hours > 0:
                    detail_data.append({
                        '教師姓名': teacher.name,
                        '行政職位': teacher.position,
                        '課程代碼': record.training.course_code,
                        '研習課程名稱': record.training.title,
                        '主辦單位': record.training.organizer,
                        '實核時數': record.hours,
                        '通過狀態': record.is_passed
                    })
        df_detail = pd.DataFrame(detail_data)
        
        # 使用 pandas 的 ExcelWriter 寫入多個工作表
        with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
            df_summary.to_excel(writer, sheet_name='全校研習統計總表', index=False)
            df_detail.to_excel(writer, sheet_name='通過研習詳細明細', index=False)
            
        return file_path

    @classmethod
    def export_official_word_report(cls, output_dir: str) -> str:
        """
        八、報表輸出功能：匯出正式簽呈公文格式 (Word .docx)
        自動建立公文文號、主旨、說明與統計表格。
        """
        file_path = os.path.join(output_dir, '教師研習檢核簽呈公文.docx')
        
        doc = Document()
        
        # 1. 設定公文大標題 (置中、粗體、24pt)
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run('臺北市大安區幸安國民小學  函 (稿)')
        title_run.bold = True
        title_run.font.size = Pt(24)
        
        # 2. 基本發文資訊
        info_p = doc.add_paragraph()
        info_p.add_run('受文者：臺北市政府教育局\n')
        info_p.add_run('發文日期：中華民國115年06月23日\n')
        info_p.add_run('字號：幸安小字第1150000123號\n')
        info_p.add_run('密等及解密條件或保密期限：普通\n')
        info_p.add_run('附件：全校教師研習統計表CSV檔')
        
        # 3. 公文結構：主旨、說明
        doc.add_heading('主旨：', level=2)
        p_subject = doc.add_paragraph('謹檢送本校「114學年度推動數位學習精進方案」全校教師研習檢核統計結果與審查報表，請鑒核。')
        
        doc.add_heading('說明：', level=2)
        p_desc = doc.add_paragraph()
        p_desc.add_run('一、依據貴局推動教師數位學習工作坊認證增能計畫辦理。\n')
        p_desc.add_run('二、本校已透過「教師研習檢核與統計管理系統 (TTVS)」完成全校教師紀錄逐筆檢核，剔除時數未達標準者。\n')
        p_desc.add_run('三、本次檢核重點包含 [精進數位] 與 [數位學習工作坊] 系列之 A1 至 B5-2 完成情形，統計摘要如下表：')
        
        # 4. 在 Word 中動態產生統計摘要表格
        df_summary = StatisticsService.generate_all_school_report()
        
        # 建立 Word 表格 (資料列數 + 1 標頭, 4 欄)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # 設定標頭
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '序號'
        hdr_cells[1].text = '行政職位'
        hdr_cells[2].text = '教師姓名'
        hdr_cells[3].text = '已通過政策標籤'
        
        # 填入前 10 筆教師資料作為公文主體內文 (其餘可見外部附件 CSV)
        for idx, row in df_summary.head(10).iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['序號'])
            row_cells[1].text = str(row['職位'])
            row_cells[2].text = str(row['姓名'])
            row_cells[3].text = str(row['通過研習'])
            
        # 5. 簽核欄位
        doc.add_paragraph('\n\n')
        sign_p = doc.add_paragraph()
        sign_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sign_p.add_run('校長  O O O  (簽章)      ')
        
        # 儲存 Word 檔案
        doc.save(file_path)
        return file_path